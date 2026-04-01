"""Word document parser (.docx, .doc)."""

from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table


def parse_word(file_path: Path) -> list[dict[str, Any]]:
    """
    Parse Word documents, preserving heading hierarchy and tables.

    Args:
        file_path: Path to the Word file.

    Returns:
        List of chunks with text and metadata.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Failed to open Word document: {e}") from e

    chunks = []
    current_section = []
    current_heading = "Document"
    section_index = 1

    for element in doc.element.body:
        if element.tag.endswith("p"):
            # Paragraph element
            para = _find_paragraph(doc, element)
            if para is not None:
                style_name = para.style.name if para.style else ""

                if style_name.startswith("Heading"):
                    # Save previous section
                    if current_section:
                        text = "\n".join(current_section)
                        if text.strip():
                            chunks.append(
                                {
                                    "text": f"{current_heading}\n\n{text}",
                                    "metadata": {
                                        "source": file_path.name,
                                        "file_type": "word",
                                        "page_or_sheet_or_slide": section_index,
                                        "heading": current_heading,
                                    },
                                }
                            )
                            section_index += 1

                    current_heading = para.text or "Section"
                    current_section = []
                else:
                    if para.text.strip():
                        current_section.append(para.text)

        elif element.tag.endswith("tbl"):
            # Table element
            table = _find_table(doc, element)
            if table is not None:
                table_text = _extract_table_text(table)
                if table_text:
                    current_section.append(table_text)

    # Save last section
    if current_section:
        text = "\n".join(current_section)
        if text.strip():
            chunks.append(
                {
                    "text": f"{current_heading}\n\n{text}",
                    "metadata": {
                        "source": file_path.name,
                        "file_type": "word",
                        "page_or_sheet_or_slide": section_index,
                        "heading": current_heading,
                    },
                }
            )

    return chunks


def _find_paragraph(doc: Document, element) -> Any:
    """Find paragraph object matching element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc: Document, element) -> Table | None:
    """Find table object matching element."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _extract_table_text(table: Table) -> str:
    """Extract text from a Word table as structured text."""
    rows = []

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))

    if not rows:
        return ""

    # First row as header
    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []

    result = f"Table:\n{header}\n"
    result += "-" * len(header) + "\n"
    result += "\n".join(body)

    return result
